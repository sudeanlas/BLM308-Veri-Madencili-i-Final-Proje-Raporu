from pathlib import Path
from math import erfc, sqrt

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.data import load_dataset
from src.modeling import train_one_model, evaluate_pipeline
from src.localization import column_label


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "teslim"
FIG = ROOT / "figures"
OUT.mkdir(exist_ok=True)

SCENARIO_LABELS = {
    "duration dahil": "Görüşme süresi dahil",
    "duration hariç": "Görüşme süresi hariç",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(text))
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=(255, 255, 255))
        set_cell_shading(table.rows[0].cells[i], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_para(doc, text, style=None, bold_start=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.space_after = Pt(6)
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_picture_if_exists(doc, path, width=5.8, caption=None):
    path = Path(path)
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            c = doc.add_paragraph(caption)
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.runs[0].italic = True
            c.paragraph_format.space_after = Pt(9)


def make_turkish_target_figure(df):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import seaborn as sns

    tmp = df.copy()
    tmp["hedef"] = tmp["y"].map({"yes": "Abone oldu", "no": "Abone olmadı"})
    plt.figure(figsize=(7, 4))
    ax = sns.countplot(data=tmp, x="hedef", hue="hedef", palette=["#9aa3af", "#2563eb"], legend=False)
    ax.set_title("Hedef değişken dağılımı")
    ax.set_xlabel("Vadeli mevduat aboneliği")
    ax.set_ylabel("Müşteri sayısı")
    plt.tight_layout()
    out = FIG / "target_distribution_turkce.png"
    plt.savefig(out, dpi=180)
    plt.close()
    return out


def make_turkish_correlation_figure(df):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import seaborn as sns

    numeric = df.select_dtypes(exclude=["object", "string"]).copy()
    numeric = numeric.rename(columns={col: column_label(col) for col in numeric.columns})
    plt.figure(figsize=(8.2, 6.2))
    sns.heatmap(numeric.corr(), cmap="Blues", center=0, linewidths=.4)
    plt.title("Sayısal değişkenler korelasyon matrisi")
    plt.tight_layout()
    out = FIG / "correlation_matrix_turkce.png"
    plt.savefig(out, dpi=180)
    plt.close()
    return out


def make_turkish_auc_figure(metrics):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    data = metrics.copy()
    data["Senaryo"] = data["Senaryo"].map(SCENARIO_LABELS).fillna(data["Senaryo"])
    pivot = data.pivot(index="Model", columns="Senaryo", values="ROC-AUC")
    plt.figure(figsize=(8, 4.5))
    pivot.plot(kind="bar", ax=plt.gca(), color=["#2563eb", "#f97316"])
    plt.title("Modellerin ROC-AUC karşılaştırması")
    plt.ylabel("ROC-AUC")
    plt.xlabel("Model")
    plt.ylim(0.5, 1.0)
    plt.xticks(rotation=20, ha="right")
    plt.tight_layout()
    out = FIG / "model_auc_comparison_turkce.png"
    plt.savefig(out, dpi=180)
    plt.close()
    return out


def mcnemar_rf_vs_mlp():
    df = load_dataset()
    rf, _, X_test, _, y_test = train_one_model(df, "Random Forest", include_duration=False)
    mlp, _, _, _, _ = train_one_model(df, "MLP", include_duration=False)
    rf_correct = rf.predict(X_test) == y_test
    mlp_correct = mlp.predict(X_test) == y_test
    b = int((rf_correct & ~mlp_correct).sum())
    c = int((~rf_correct & mlp_correct).sum())
    stat = ((abs(b - c) - 1) ** 2) / (b + c) if (b + c) else 0
    # For chi-square with 1 degree of freedom, survival function is erfc(sqrt(x/2)).
    p_value = float(erfc(sqrt(stat / 2))) if stat > 0 else 1.0
    return b, c, stat, p_value


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.color.rgb = RGBColor(31, 78, 121)


def build_docx():
    metrics = pd.read_csv(OUT / "metrics.csv")
    b, c, stat, p_value = mcnemar_rf_vs_mlp()
    df = load_dataset()
    target_fig = make_turkish_target_figure(df)
    corr_fig = make_turkish_correlation_figure(df)
    auc_fig = make_turkish_auc_figure(metrics)
    yes_count = int((df["y"] == "yes").sum())
    no_count = int((df["y"] == "no").sum())

    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("BLM308 Veri Madenciliği Final Proje Raporu")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run("Banka Pazarlama Kampanyalarında Müşteri Tahmin ve Açıklanabilirlik Sistemi")
    s.bold = True
    s.font.size = Pt(15)
    s.font.color.rgb = RGBColor(31, 78, 121)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Proje türü: Interpretable ML / XAI\nTakım büyüklüğü: 3 kişi\nBahar 2026")
    doc.add_paragraph()

    add_table(doc, ["Ad Soyad", "Numara", "Rol", "Sorumlu Olduğu Bölümler"], [
        ["Sude ANLAŞ", "231041038", "Veri Mühendisi", "Veri setinin hazırlanması, veri anlama, EDA, ön işleme, veri sızıntısı analizi"],
        ["Doğan Fatih OĞUR", "231041048", "Model Mühendisi", "Sınıflandırma modelleri, Random Forest/MLP kurulumu, Streamlit arayüzü, kod düzeni"],
        ["Merve Naz ORAN", "231041058", "Değerlendirme ve XAI Analisti", "Metrik karşılaştırması, McNemar testi, SHAP/LIME yorumları, sonuç ve iş değeri analizi"],
    ], widths=[1.4, 1.0, 1.9, 3.8])

    add_para(doc, "Proje türü: Interpretable ML / XAI")
    add_para(doc, "Takım büyüklüğüne göre ek kapsam: 3 kişilik proje beklentisine uygun olarak temel sınıflandırıcı karşılaştırmasına ek olarak Random Forest ve MLP kara kutu modelleri, açıklanabilirlik analizi, Streamlit tabanlı uygulama arayüzü ve McNemar istatistiksel karşılaştırması eklenmiştir.")
    doc.add_page_break()

    doc.add_heading("Özet", level=1)
    add_para(doc, "Bu çalışmada UCI Bank Marketing veri seti kullanılarak banka pazarlama kampanyalarında bir müşterinin vadeli mevduata abone olma ihtimali tahmin edilmiştir. Proje, CRISP-DM metodolojisinin iş anlayışı, veri anlayışı, veri hazırlığı, modelleme, değerlendirme ve iş değeri adımlarına göre yürütülmüştür. Lojistik Regresyon ve Karar Ağacı temel karşılaştırma modelleri olarak; Random Forest ve Çok Katmanlı Algılayıcı ise kara kutu model olarak kullanılmıştır. Model başarısı doğruluk, kesinlik, duyarlılık, F1 ve ROC-AUC metrikleriyle değerlendirilmiştir. Ayrıca telefon görüşmesi tamamlandıktan sonra bilinen görüşme süresi değişkeninin veri sızıntısı riski tartışılmış ve görüşme süresi dahil/hariç iki senaryo karşılaştırılmıştır. Random Forest modelinin kararları değişken önem düzeyleri, SHAP/LIME yaklaşımı ve yerel açıklama mantığıyla yorumlanmıştır. Üç kişilik takım kapsamına uygun olarak, Random Forest ve Çok Katmanlı Algılayıcı modelleri arasında McNemar testi uygulanmış, sonuçların yalnızca sayısal metriklerden değil istatistiksel anlamlılık açısından da değerlendirilmesi sağlanmıştır. Sonuçlar, açıklanabilir makine öğrenmesi yaklaşımının bankacılık gibi güven ve gerekçelendirme gerektiren alanlarda tahmin başarısını iş değeriyle birleştirdiğini göstermektedir.")
    add_para(doc, "Anahtar kelimeler: veri madenciliği, CRISP-DM, sınıflandırma, Random Forest, MLP, SHAP, LIME, açıklanabilir yapay zeka, bankacılık pazarlaması.")

    doc.add_heading("1. Giriş", level=1)
    doc.add_heading("1.1 Problem Tanımı", level=2)
    add_para(doc, "Bankalar pazarlama kampanyalarında çağrı merkezi kaynaklarını verimli kullanmak zorundadır. Her müşteriyi aynı öncelikle aramak, maliyeti artırabilir ve kampanyadan alınan dönüşüm oranını düşürebilir. Bu projede temel problem, geçmiş müşteri, kampanya ve ekonomik gösterge verilerinden yararlanarak bir müşterinin vadeli mevduata abone olup olmayacağını tahmin etmektir.")
    add_para(doc, "Bu problem yalnızca sınıflandırma başarısı açısından ele alınmamıştır. Bankacılık alanında modelin neden belirli bir müşteriyi yüksek veya düşük olasılıklı gördüğünün de açıklanması gerekir. Bu nedenle proje, yönergede belirtilen Interpretable ML / XAI başlığına doğrudan karşılık gelecek biçimde tasarlanmıştır.")
    doc.add_heading("1.2 Motivasyon ve İş Değeri", level=2)
    add_para(doc, "Tahmin modeli, banka pazarlama ekibinin yüksek dönüşüm potansiyeli olan müşterilere öncelik vermesine yardımcı olabilir. Örneğin arama listesi model skoruna göre sıralandığında kampanya ekibi sınırlı zamanını daha verimli kullanabilir. Açıklanabilirlik katmanı ise pazarlama uzmanının model önerisini körü körüne kabul etmek yerine karar gerekçesini anlamasını sağlar.")
    add_para(doc, "Model çıktıları pratikte üç şekilde değer üretir: çağrı önceliklendirme, kampanya segmentasyonu ve yönetilebilir karar açıklaması. Böylece proje, yalnızca teknik bir sınıflandırma çalışması değil, iş problemine uygulanabilir bir karar destek yaklaşımıdır.")
    doc.add_heading("1.3 Literatür ve İlgili Çalışmalar", level=2)
    add_para(doc, "Bank Marketing veri seti, Moro, Cortez ve Rita tarafından bankacılık telepazarlama başarısını tahmin etmek için kullanılan çalışmayla ilişkilidir. Veri seti, müşteri özellikleri, kampanya geçmişi ve ekonomik göstergeler üzerinden ikili sınıflandırma yapılmasına olanak tanır.")
    add_para(doc, "Açıklanabilir yapay zeka alanında SHAP, Shapley değerlerine dayalı olarak her değişkenin tahmine katkısını yorumlamayı sağlar. LIME ise belirli bir örnek etrafında yerel ve anlaşılır bir model kurarak kara kutu model kararlarını açıklar. Bu projede bu yaklaşımların mantığı, Random Forest model kararlarının yorumlanmasında kullanılmıştır.")
    add_para(doc, "Ders kapsamında işlenen veri madenciliği süreci dikkate alınarak rapor, yalnızca algoritma sonuçlarını değil veri hazırlama, model seçimi, değerlendirme, istatistiksel karşılaştırma ve iş değeri yorumunu da içerecek şekilde yapılandırılmıştır.")
    doc.add_heading("1.4 Raporun Yapısı", level=2)
    add_para(doc, "İkinci bölümde CRISP-DM süreci ve kullanılan araçlar, üçüncü bölümde veri seti ve ön işleme adımları, dördüncü bölümde modelleme yaklaşımı, beşinci bölümde sonuçlar ve istatistiksel değerlendirme, altıncı bölümde açıklanabilirlik, yedinci bölümde uygulama arayüzü, son bölümde ise sonuçlar ve gelecek çalışmalar sunulmuştur.")

    doc.add_heading("2. Yöntem: CRISP-DM", level=1)
    add_para(doc, "Proje, veri madenciliği projelerinde yaygın olarak kullanılan CRISP-DM metodolojisine göre yürütülmüştür. Bu yöntem, projenin yalnızca model eğitimiyle sınırlı kalmamasını, iş probleminden son yorumlara kadar uçtan uca ele alınmasını sağlar.")
    add_table(doc, ["CRISP-DM Aşaması", "Bu Projedeki Karşılığı"], [
        ["İş anlayışı", "Banka pazarlama kampanyalarında doğru müşteriyi önceliklendirme problemi tanımlandı."],
        ["Veri anlayışı", "Bank Marketing veri setindeki müşteri, kampanya ve ekonomik gösterge değişkenleri incelendi."],
        ["Veri hazırlığı", "Kategorik değişkenler kodlandı, sayısal değişkenler ölçeklendirildi, görüşme süresi sızıntı riski ayrı ele alındı."],
        ["Modelleme", "Logistic Regression, Decision Tree, Random Forest ve MLP modelleri kuruldu."],
        ["Değerlendirme", "Doğruluk, kesinlik, duyarlılık, F1, ROC-AUC ve McNemar testiyle sonuçlar incelendi."],
        ["Dağıtım/yorum", "Streamlit arayüzüyle tahmin sistemi hazırlandı ve XAI çıktıları iş değeri açısından yorumlandı."],
    ], widths=[2.0, 5.8])
    doc.add_heading("2.1 Kullanılan Araçlar", level=2)
    add_table(doc, ["Araç", "Kullanım Amacı"], [
        ["Python", "Veri hazırlama, modelleme, metrik üretimi ve rapor çıktılarının hazırlanması"],
        ["scikit-learn", "Sınıflandırma modelleri, ön işleme boru hattı ve metrik hesaplama"],
        ["pandas / numpy", "Veri okuma, dönüştürme ve analiz işlemleri"],
        ["matplotlib / seaborn", "EDA, metrik ve açıklanabilirlik grafiklerinin oluşturulması"],
        ["LIME / SHAP yaklaşımı", "Kara kutu model kararlarının genel ve yerel düzeyde açıklanması"],
        ["Streamlit", "Hocaya gösterilebilecek Türkçe tahmin ve açıklanabilirlik arayüzü"],
    ], widths=[1.8, 6.0])

    doc.add_heading("3. Veri", level=1)
    doc.add_heading("3.1 Veri Kaynağı", level=2)
    add_para(doc, "Kullanılan veri seti UCI Machine Learning Repository üzerinde yayımlanan Bank Marketing veri setidir. Veri, Portekizli bir bankanın telefonla pazarlama kampanyalarına aittir. Hedef değişken, müşterinin vadeli mevduata abone olup olmadığını gösteren ikili sınıftır.")
    add_table(doc, ["Özellik", "Değer"], [
        ["Örnek sayısı", f"{len(df):,}".replace(",", ".")],
        ["Girdi değişkeni sayısı", "20"],
        ["Hedef değişken", "Vadeli mevduata abone olma durumu"],
        ["Pozitif sınıf", f"Abone oldu = {yes_count:,} (%{yes_count / len(df) * 100:.1f})".replace(",", ".")],
        ["Negatif sınıf", f"Abone olmadı = {no_count:,} (%{no_count / len(df) * 100:.1f})".replace(",", ".")],
    ], widths=[2.5, 5.0])
    add_picture_if_exists(doc, target_fig, 5.7, "Şekil 1. Hedef değişken dağılımı")

    doc.add_heading("3.2 Değişken Grupları", level=2)
    add_table(doc, ["Grup", "Örnek Değişkenler", "Açıklama"], [
        ["Müşteri profili", "Yaş, meslek, medeni durum, eğitim", "Müşterinin demografik ve sosyo-ekonomik özellikleri"],
        ["Finansal durum", "Konut kredisi, ihtiyaç kredisi, temerrüt", "Müşterinin bankacılık ürünleriyle ilişkili bilgileri"],
        ["Kampanya geçmişi", "İletişim kanalı, arama sayısı, önceki kampanya sonucu", "Mevcut ve önceki kampanya temasları"],
        ["Ekonomik göstergeler", "Piyasa faiz göstergesi, tüketici güven endeksi, istihdam göstergesi", "Kampanya dönemindeki makroekonomik bağlam"],
    ], widths=[1.7, 2.5, 3.5])
    add_picture_if_exists(doc, corr_fig, 5.8, "Şekil 2. Sayısal değişkenler korelasyon matrisi")

    doc.add_heading("3.3 Keşifsel Veri Analizi", level=2)
    add_para(doc, "Hedef değişken dağılımı, veri setinde sınıf dengesizliği olduğunu göstermektedir. Abone olmayan müşteriler veri setinin büyük çoğunluğunu oluşturur. Bu nedenle doğruluk tek başına yeterli bir değerlendirme metriği değildir; kesinlik, duyarlılık, F1 ve ROC-AUC birlikte değerlendirilmiştir.")
    add_para(doc, "Sayısal değişkenler arasındaki korelasyon incelendiğinde ekonomik göstergelerin birbirleriyle ilişkili olduğu görülmektedir. Bu durum bankacılık kampanyalarında makroekonomik koşulların müşteri davranışını etkileyebileceğini düşündürür. Ancak korelasyon nedensellik anlamına gelmediği için yorumlar model açıklanabilirliğiyle birlikte ele alınmıştır.")
    doc.add_heading("3.4 Veri Hazırlığı ve Ön İşleme", level=2)
    add_numbered(doc, [
        "Hedef değişken ikili sınıfa dönüştürülmüştür: abone oldu = 1, abone olmadı = 0.",
        "Kategorik değişkenler One-Hot Encoding yöntemiyle sayısal forma çevrilmiştir.",
        "Sayısal değişkenler StandardScaler ile ölçeklendirilmiştir. Bu işlem özellikle MLP için gereklidir.",
        "Veri stratified şekilde yüzde 75 eğitim ve yüzde 25 test olarak ayrılmıştır.",
        "Sınıf dengesizliği nedeniyle bazı modellerde class_weight kullanılmıştır.",
        "Görüşme süresi değişkeni veri sızıntısı riski taşıdığı için ayrı senaryo olarak değerlendirilmiştir.",
    ])

    doc.add_heading("4. Modelleme", level=1)
    add_para(doc, "Model seçimi, hem ders kapsamında görülen klasik sınıflandırma mantığını hem de proje konusundaki açıklanabilir kara kutu model ihtiyacını karşılayacak şekilde yapılmıştır.")
    add_table(doc, ["Model", "Tür", "Weka Karşılığı", "Projedeki Rolü"], [
        ["Logistic Regression", "Lineer temel model", "Logistic", "Basit ve yorumlanabilir referans model"],
        ["Decision Tree", "Ağaç tabanlı model", "J48", "Kural benzeri açıklanabilir karşılaştırma modeli"],
        ["Random Forest", "Ensemble kara kutu model", "trees.RandomForest", "Ana performans modeli ve açıklanabilirlik analizi"],
        ["MLP", "Sinir ağı kara kutu model", "functions.MultilayerPerceptron", "Alternatif kara kutu model"],
    ], widths=[1.5, 1.8, 1.8, 2.8])
    doc.add_heading("4.1 Deney Kurulumu", level=2)
    add_bullets(doc, [
        "Veri bölme oranı: yüzde 75 eğitim, yüzde 25 test.",
        "Bölme yöntemi: hedef değişken dağılımını koruyan stratified split.",
        "Rassallık kontrolü: random_state = 42.",
        "Ana değerlendirme metrikleri: Accuracy, Precision, Recall, F1 ve ROC-AUC.",
        "İstatistiksel karşılaştırma: Random Forest ve MLP modelleri için McNemar testi.",
        "Uygulama modeli: gerçekçi kullanım için görüşme süresi hariç Random Forest.",
    ])
    doc.add_heading("4.2 Hiperparametre Tercihleri", level=2)
    add_table(doc, ["Model", "Önemli Parametreler", "Gerekçe"], [
        ["Random Forest", "180 ağaç, max_depth=12, min_samples_leaf=5", "Varyansı düşürmek ve aşırı öğrenmeyi sınırlamak"],
        ["MLP", "48 ve 24 nöronlu iki gizli katman, early stopping", "Doğrusal olmayan ilişkileri öğrenirken aşırı öğrenmeyi azaltmak"],
        ["Decision Tree", "max_depth=8", "Ağacın aşırı derinleşmesini önlemek"],
        ["Logistic Regression", "class_weight=balanced", "Sınıf dengesizliğinin etkisini azaltmak"],
    ], widths=[1.6, 3.0, 3.2])

    doc.add_heading("5. Sonuçlar ve Değerlendirme", level=1)
    doc.add_heading("5.1 Performans Karşılaştırması", level=2)
    rows = []
    for _, row in metrics.iterrows():
        rows.append([
            SCENARIO_LABELS.get(row["Senaryo"], row["Senaryo"]), row["Model"],
            f"{row['Accuracy']:.3f}", f"{row['Precision']:.3f}",
            f"{row['Recall']:.3f}", f"{row['F1']:.3f}", f"{row['ROC-AUC']:.3f}",
        ])
    add_table(doc, ["Senaryo", "Model", "Doğruluk", "Kesinlik", "Duyarlılık", "F1", "ROC-AUC"], rows, widths=[1.45, 1.45, .75, .75, .85, .55, .75])
    add_picture_if_exists(doc, auc_fig, 5.8, "Şekil 3. Modellerin ROC-AUC karşılaştırması")
    add_para(doc, "Görüşme süresi dahil senaryoda Random Forest en yüksek ROC-AUC değerine ulaşmıştır. Ancak bu sonuç gerçek kampanya öncesi kullanım için doğrudan tercih edilmemelidir; çünkü görüşme süresi değişkeni görüşme tamamlandıktan sonra oluşur. Gerçekçi kullanımda görüşme süresi hariç senaryo esas alınmalıdır.")
    add_para(doc, "Görüşme süresi hariç senaryoda Random Forest, ROC-AUC açısından güçlü ve dengeli bir sonuç vermiştir. Çok Katmanlı Algılayıcı modeli doğruluk açısından yüksek görünse de pozitif sınıf duyarlılık değeri düşüktür. Banka pazarlaması probleminde yalnızca abone olmayacak müşterileri doğru tahmin etmek yeterli değildir; abone olabilecek müşterileri kaçırmamak da önemlidir.")

    doc.add_heading("5.2 Confusion Matrix Yorumu", level=2)
    add_picture_if_exists(doc, FIG / "rf_confusion_matrix.png", 4.9, "Şekil 4. Random Forest confusion matrix")
    add_para(doc, "Hata matrisi, modelin abone olmayan müşteri sınıfında güçlü olduğunu, abone olabilecek müşteri sınıfında ise sınıf dengesizliği nedeniyle daha dikkatli yorumlanması gerektiğini gösterir. Bankacılık kampanyasında yanlış negatif, potansiyel olarak kampanyaya olumlu dönebilecek bir müşterinin düşük öncelikli görülmesi anlamına gelir. Bu nedenle duyarlılık değeri iş açısından önemlidir.")

    doc.add_heading("5.3 İstatistiksel Test: McNemar Karşılaştırması", level=2)
    add_para(doc, "Üç kişilik takım kapsamına uygun olarak değerlendirme bölümüne istatistiksel karşılaştırma eklenmiştir. Random Forest ve Çok Katmanlı Algılayıcı modellerinin görüşme süresi hariç gerçekçi senaryodaki test tahminleri McNemar testiyle karşılaştırılmıştır. McNemar testi, aynı test örnekleri üzerinde iki sınıflandırıcının hata örüntülerinin anlamlı biçimde farklı olup olmadığını ölçer.")
    add_table(doc, ["Karşılaştırma", "b", "c", "Test istatistiği", "p-değeri", "Yorum"], [[
        "Random Forest vs MLP", b, c, f"{stat:.3f}", f"{p_value:.4f}",
        "p < 0.05 ise iki modelin hata örüntüsü istatistiksel olarak farklı kabul edilir." if p_value < 0.05 else "p >= 0.05 olduğundan anlamlı fark yorumu temkinli yapılmalıdır.",
    ]], widths=[1.9, .55, .55, 1.2, 1.0, 3.2])
    add_para(doc, "Bu test, yalnızca metrik tablosuna bakmak yerine modellerin aynı müşteriler üzerindeki doğru/yanlış kararlarının karşılaştırılmasını sağlar. Böylece değerlendirme bölümü, yönergede vurgulanan metrik ve test beklentisini daha güçlü şekilde karşılar.")

    doc.add_heading("5.4 Bias-Variance Yorumu", level=2)
    add_para(doc, "Logistic Regression düşük varyanslı ancak doğrusal varsayımları nedeniyle sınırlı bir modeldir. Decision Tree yorumlanabilir olmakla birlikte tekil ağaç yapısı nedeniyle yüksek varyansa yatkındır. Random Forest çok sayıda karar ağacını birleştirdiği için varyansı azaltır ve daha kararlı tahminler üretir. MLP ise doğrusal olmayan ilişkileri öğrenebilir; ancak eğitim süreci, ölçekleme ve düzenlileştirme yapılmadığında aşırı öğrenmeye yatkın olabilir.")
    add_para(doc, "Bu proje için Random Forest, performans, kararlılık ve açıklanabilirlik araçlarıyla uyumluluk açısından en uygun model olarak değerlendirilmiştir. MLP alternatif kara kutu model olarak yararlı bir karşılaştırma sağlamıştır.")

    doc.add_heading("6. Açıklanabilir Yapay Zeka", level=1)
    add_para(doc, "Kara kutu modeller yüksek tahmin başarısı sağlayabilir; ancak karar gerekçesi anlaşılmadığında özellikle finans alanında güven problemi oluşur. Bu nedenle model açıklanabilirliği proje konusunun merkezinde yer almaktadır.")
    doc.add_heading("6.1 Global Açıklama", level=2)
    add_picture_if_exists(doc, FIG / "rf_feature_importance_from_explain.png", 5.9, "Şekil 5. Modelin genel kararlarını etkileyen faktörler")
    add_para(doc, "Global açıklama, modelin genel olarak hangi değişkenlere daha fazla ağırlık verdiğini gösterir. Örneğin önceki kampanya sonucu, ekonomik göstergeler ve kampanya temas sayıları modelin kararlarında önemli rol oynayabilir. Bu tür bilgiler pazarlama ekibinin yalnızca tekil tahminleri değil, müşteri segmentlerini de anlamasına yardımcı olur.")
    doc.add_heading("6.2 Yerel Açıklama", level=2)
    add_picture_if_exists(doc, FIG / "lime_explanation_from_explain.png", 5.9, "Şekil 6. Tek müşteri için yerel karar açıklaması")
    add_para(doc, "Yerel açıklama, belirli bir müşteri için modelin kararına hangi koşulların katkı verdiğini gösterir. Böylece modelin bir müşteriye neden yüksek veya düşük abonelik olasılığı verdiği daha anlaşılır hale gelir. Bu açıklama doğrudan nedensellik göstermez; ancak karar destek sistemi olarak modelin şeffaflığını artırır.")
    doc.add_heading("6.3 SHAP ve LIME Yaklaşımının Projedeki Yeri", level=2)
    add_para(doc, "SHAP yaklaşımı, değişken katkılarını oyun teorisinden gelen Shapley değerleriyle yorumlamaya dayanır. LIME ise belirli bir örnek çevresinde basit ve anlaşılır bir yerel model kurar. Projede bu iki yaklaşım, Random Forest gibi karmaşık model kararlarının hem genel hem de müşteri bazında açıklanması için kullanılmıştır.")
    add_para(doc, "Açıklanabilirlik çıktıları, banka çalışanının model kararını denetlemesine ve iş bağlamında sorgulamasına olanak tanır. Bu yönüyle proje, yönergede istenen Interpretable ML / XAI kapsamını yalnızca teorik değil uygulamalı olarak da karşılar.")

    doc.add_heading("7. Uygulama Arayüzü", level=1)
    add_para(doc, "Projede rapor ve kod çıktılarının yanında, hocaya gösterilebilecek Türkçe bir tahmin arayüzü de hazırlanmıştır. Arayüz Streamlit ile geliştirilmiştir ve Visual Studio Code üzerinden çalıştırılabilir.")
    add_bullets(doc, [
        "Sol bölümde müşteri bilgileri Türkçe değişken adlarıyla girilir.",
        "Sağ bölümde modelin vadeli mevduata abone olma olasılığı gösterilir.",
        "Alt bölümde modelin genel kararlarını etkileyen faktörler Türkçe adlarla listelenir.",
        "Arayüzde teknik olmayan İngilizce ifadeler kullanılmaz; müşteri alanları ve sonuçlar Türkçe gösterilir.",
    ])
    add_para(doc, "Uygulama gerçekçi kampanya öncesi senaryo için görüşme süresi değişkenini kullanmaz. Bu tercih, raporda tartışılan veri sızıntısı riskini arayüz tarafında da dikkate aldığımızı gösterir.")
    add_table(doc, ["Komut", "Açıklama"], [
        ["python train.py", "Modelleri eğitir, metrikleri üretir ve Random Forest modelini kaydeder."],
        ["python explain.py", "Açıklanabilirlik grafiklerini ve tablo çıktılarını üretir."],
        ["python -m streamlit run app.py", "Türkçe müşteri tahmin ve açıklanabilirlik arayüzünü açar."],
    ], widths=[3.0, 4.7])

    doc.add_heading("8. İş Değeri, Etik ve Sınırlılıklar", level=1)
    doc.add_heading("8.1 İş Değeri", level=2)
    add_para(doc, "Model, kampanya listesinde müşterileri abone olma olasılığına göre sıralamak için kullanılabilir. Böylece çağrı merkezi ekibi yüksek potansiyelli müşterilere daha erken ulaşabilir. Bu yaklaşım kampanya maliyetini düşürebilir, dönüşüm oranını artırabilir ve müşteri iletişim stratejisini daha veriye dayalı hale getirebilir.")
    doc.add_heading("8.2 Etik ve Veri Koruma", level=2)
    add_para(doc, "Finansal hizmetlerde müşteri verilerinin kullanımı dikkatli yönetilmelidir. Model çıktıları doğrudan müşteri aleyhine otomatik karar olarak kullanılmamalı, yalnızca pazarlama önceliklendirme ve karar destek amacıyla değerlendirilmelidir. Kişisel verilerin işlenmesinde KVKK/GDPR benzeri veri koruma ilkeleri dikkate alınmalıdır.")
    doc.add_heading("8.3 Sınırlılıklar", level=2)
    add_bullets(doc, [
        "Veri seti belirli bir ülke ve dönem kampanyalarına ait olduğu için genellenebilirlik sınırlıdır.",
        "Abone olabilecek müşteri sınıfı azınlıkta olduğu için duyarlılık ve kesinlik dengesi dikkatli yorumlanmalıdır.",
        "Görüşme süresi değişkeni kampanya öncesi kullanımda veri sızıntısı oluşturabileceği için ana uygulamada kullanılmamıştır.",
        "Model açıklamaları nedensellik göstermez; yalnızca tahmin davranışını yorumlamaya yardım eder.",
        "Gerçek banka ortamında model periyodik olarak izlenmeli ve yeniden eğitilmelidir.",
    ])

    doc.add_heading("9. Sonuç ve Gelecek Çalışmalar", level=1)
    add_para(doc, "Bu proje, Bank Marketing veri seti üzerinde müşterinin vadeli mevduata abone olma ihtimalini tahmin eden ve model kararlarını açıklanabilir hale getiren uçtan uca bir veri madenciliği çalışmasıdır. Logistic Regression, Decision Tree, Random Forest ve MLP modelleri karşılaştırılmış; Random Forest performans ve açıklanabilirlik dengesi açısından öne çıkmıştır.")
    add_para(doc, "Üç kişilik takım kapsamına uygun olarak projeye yalnızca XAI değil, uygulama arayüzü ve McNemar istatistiksel karşılaştırması da eklenmiştir. Böylece proje, yönergede beklenen modelleme, değerlendirme, iş değeri, rol dağılımı ve rapor kalitesi kriterlerini kapsamlı biçimde karşılamaktadır.")
    add_para(doc, "Gelecek çalışmalarda maliyet duyarlı öğrenme, karar eşiği optimizasyonu, zaman bazlı validasyon, farklı XAI yöntemlerinin karşılaştırılması ve gerçek banka kampanyalarına daha yakın yerel veriyle test yapılması önerilebilir.")

    doc.add_heading("Kaynaklar", level=1)
    refs = [
        "Moro, S., Cortez, P., Rita, P. (2014). A Data-Driven Approach to Predict the Success of Bank Telemarketing. Decision Support Systems.",
        "UCI Machine Learning Repository. Bank Marketing Dataset. https://archive.ics.uci.edu/dataset/222/bank+marketing",
        "Lundberg, S. M., Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions. NeurIPS.",
        "Ribeiro, M. T., Singh, S., Guestrin, C. (2016). Why Should I Trust You? Explaining the Predictions of Any Classifier. NAACL-HLT.",
        "Witten, I. H., Frank, E., Hall, M. A., Pal, C. J. Data Mining: Practical Machine Learning Tools and Techniques.",
        "Han, J., Kamber, M., Pei, J. Data Mining: Concepts and Techniques.",
        "scikit-learn documentation. https://scikit-learn.org",
        "Streamlit documentation. https://streamlit.io",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    doc.add_heading("Ek A: Teslim Dosya Yapısı", level=1)
    add_para(doc, "Teslim paketinde rapor.docx, rapor.tex, sunum_beamer.tex, README dosyaları, Python kodları, veri dosyası, eğitilmiş model, grafikler ve Streamlit arayüzü bulunmaktadır.")
    add_table(doc, ["Klasör/Dosya", "İçerik"], [
        ["teslim/rapor.docx", "Ana Word raporu"],
        ["train.py", "Model eğitimi ve metrik üretimi"],
        ["explain.py", "Açıklanabilirlik grafikleri"],
        ["app.py", "Türkçe tahmin ve açıklanabilirlik arayüzü"],
        ["src/", "Veri yükleme, modelleme, XAI ve Türkçe etiketleme modülleri"],
        ["veri/", "Bank Marketing veri seti"],
        ["figures/", "EDA, metrik ve açıklanabilirlik grafikleri"],
        ["models/", "Eğitilmiş Random Forest modeli"],
    ], widths=[2.6, 5.1])

    doc.add_heading("Ek B: Yapay Zeka Kullanımı Beyanı", level=1)
    add_para(doc, "Bu raporun taslak oluşturma, dil düzenleme ve proje dosya organizasyonu aşamalarında yapay zeka desteği kullanılmıştır. Veri seti, modelleme kodu, metrikler, grafikler ve yorumlar proje ekibi tarafından kontrol edilerek düzenlenmiştir. Yapay zeka çıktıları doğrudan kopyalanmadan, ders yönergesine ve akademik rapor formatına uygun şekilde revize edilmiştir.")

    out_path = OUT / "rapor.docx"
    backup_path = OUT / "rapor_2_kisilik_onceki_surumu.docx"
    if out_path.exists() and not backup_path.exists():
        backup_path.write_bytes(out_path.read_bytes())
    doc.save(out_path)
    also_path = OUT / "rapor_3_kisilik_revize.docx"
    doc.save(also_path)
    print(out_path)
    print(also_path)
    print(f"McNemar: b={b}, c={c}, stat={stat:.4f}, p={p_value:.6f}")


if __name__ == "__main__":
    build_docx()

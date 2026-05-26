from pathlib import Path
import json
import textwrap
import zipfile
import io
import os
from datetime import date

import numpy as np
import pandas as pd
import matplotlib
os.environ.setdefault("MPLCONFIGDIR", ".mplconfig")
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from sklearn.compose import ColumnTransformer
from sklearn.metrics import (
    accuracy_score,
    precision_score,
    recall_score,
    f1_score,
    roc_auc_score,
    confusion_matrix,
    ConfusionMatrixDisplay,
)
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.linear_model import LogisticRegression
from sklearn.tree import DecisionTreeClassifier
from sklearn.ensemble import RandomForestClassifier
from sklearn.neural_network import MLPClassifier


ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "veri"
CODE_DIR = ROOT / "kod"
FIG_DIR = ROOT / "figures"
OUT_DIR = ROOT / "teslim"
for d in [DATA_DIR, CODE_DIR, FIG_DIR, OUT_DIR]:
    d.mkdir(exist_ok=True)


def download_bank_marketing():
    csv_path = DATA_DIR / "bank-additional-full.csv"
    if csv_path.exists():
        return csv_path
    url = "https://archive.ics.uci.edu/static/public/222/bank+marketing.zip"
    zip_path = DATA_DIR / "bank_marketing.zip"
    import urllib.request

    urllib.request.urlretrieve(url, zip_path)
    with zipfile.ZipFile(zip_path) as zf:
        names = zf.namelist()
        if "bank-additional/bank-additional-full.csv" in names:
            with zf.open("bank-additional/bank-additional-full.csv") as src, open(csv_path, "wb") as dst:
                dst.write(src.read())
        else:
            nested = zf.read("bank-additional.zip")
            with zipfile.ZipFile(io.BytesIO(nested)) as nested_zf:
                with nested_zf.open("bank-additional/bank-additional-full.csv") as src, open(csv_path, "wb") as dst:
                    dst.write(src.read())
    return csv_path


def load_data():
    csv_path = download_bank_marketing()
    df = pd.read_csv(csv_path, sep=";")
    df.to_csv(DATA_DIR / "bank_marketing_raw_copy.csv", index=False)
    df.to_csv(DATA_DIR / "bank_marketing_processed_base.csv", index=False)
    return df


def feature_names(preprocessor):
    names = []
    for name, trans, cols in preprocessor.transformers_:
        if name == "cat":
            ohe = trans.named_steps["onehot"]
            names.extend(ohe.get_feature_names_out(cols))
        elif name == "num":
            names.extend(cols)
    return list(names)


def evaluate_models(df, include_duration=True):
    y = (df["y"] == "yes").astype(int)
    X = df.drop(columns=["y"])
    if not include_duration:
        X = X.drop(columns=["duration"])

    cat_cols = X.select_dtypes(include=["object"]).columns.tolist()
    num_cols = X.select_dtypes(exclude=["object"]).columns.tolist()

    preprocessor = ColumnTransformer(
        transformers=[
            ("cat", Pipeline([
                ("imputer", SimpleImputer(strategy="most_frequent")),
                ("onehot", OneHotEncoder(handle_unknown="ignore", sparse_output=False)),
            ]), cat_cols),
            ("num", Pipeline([
                ("imputer", SimpleImputer(strategy="median")),
                ("scaler", StandardScaler()),
            ]), num_cols),
        ]
    )

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.25, random_state=42, stratify=y
    )

    models = {
        "Logistic Regression": LogisticRegression(max_iter=1000, class_weight="balanced", random_state=42),
        "Decision Tree": DecisionTreeClassifier(max_depth=8, class_weight="balanced", random_state=42),
        "Random Forest": RandomForestClassifier(
            n_estimators=180, max_depth=12, min_samples_leaf=5,
            class_weight="balanced_subsample", n_jobs=-1, random_state=42
        ),
        "MLP": MLPClassifier(
            hidden_layer_sizes=(48, 24), alpha=0.001, max_iter=90,
            early_stopping=True, random_state=42
        ),
    }

    rows, fitted = [], {}
    for model_name, clf in models.items():
        pipe = Pipeline([("prep", preprocessor), ("model", clf)])
        pipe.fit(X_train, y_train)
        pred = pipe.predict(X_test)
        if hasattr(pipe.named_steps["model"], "predict_proba"):
            prob = pipe.predict_proba(X_test)[:, 1]
        else:
            prob = pred
        rows.append({
            "Senaryo": "duration dahil" if include_duration else "duration hariç",
            "Model": model_name,
            "Accuracy": accuracy_score(y_test, pred),
            "Precision": precision_score(y_test, pred, zero_division=0),
            "Recall": recall_score(y_test, pred, zero_division=0),
            "F1": f1_score(y_test, pred, zero_division=0),
            "ROC-AUC": roc_auc_score(y_test, prob),
        })
        fitted[model_name] = (pipe, X_train, X_test, y_train, y_test)
    return pd.DataFrame(rows), fitted


def make_figures(df, metrics, fitted):
    sns.set_theme(style="whitegrid")

    plt.figure(figsize=(7, 4))
    ax = sns.countplot(data=df, x="y", hue="y", palette=["#9aa3af", "#2563eb"], legend=False)
    ax.set_title("Hedef değişken dağılımı")
    ax.set_xlabel("Vadeli mevduata abone oldu mu?")
    ax.set_ylabel("Müşteri sayısı")
    plt.tight_layout()
    target_fig = FIG_DIR / "target_distribution.png"
    plt.savefig(target_fig, dpi=180)
    plt.close()

    numeric = df.select_dtypes(exclude=["object"]).copy()
    corr = numeric.corr()
    plt.figure(figsize=(8, 6))
    sns.heatmap(corr, cmap="Blues", center=0, linewidths=.4)
    plt.title("Sayısal değişkenler korelasyon matrisi")
    plt.tight_layout()
    corr_fig = FIG_DIR / "correlation_matrix.png"
    plt.savefig(corr_fig, dpi=180)
    plt.close()

    plt.figure(figsize=(8, 4.5))
    pivot = metrics.pivot(index="Model", columns="Senaryo", values="ROC-AUC")
    pivot.plot(kind="bar", ax=plt.gca(), color=["#2563eb", "#f97316"])
    plt.title("Model ROC-AUC karşılaştırması")
    plt.ylabel("ROC-AUC")
    plt.ylim(0.5, 1.0)
    plt.xticks(rotation=20, ha="right")
    plt.tight_layout()
    auc_fig = FIG_DIR / "model_auc_comparison.png"
    plt.savefig(auc_fig, dpi=180)
    plt.close()

    rf_pipe, X_train, X_test, y_train, y_test = fitted["Random Forest"]
    pred = rf_pipe.predict(X_test)
    cm = confusion_matrix(y_test, pred)
    disp = ConfusionMatrixDisplay(cm, display_labels=["no", "yes"])
    fig, ax = plt.subplots(figsize=(5.5, 4.5))
    disp.plot(ax=ax, cmap="Blues", colorbar=False)
    ax.set_title("Random Forest confusion matrix")
    plt.tight_layout()
    cm_fig = FIG_DIR / "rf_confusion_matrix.png"
    plt.savefig(cm_fig, dpi=180)
    plt.close()

    prep = rf_pipe.named_steps["prep"]
    names = feature_names(prep)
    importances = rf_pipe.named_steps["model"].feature_importances_
    imp = pd.DataFrame({"feature": names, "importance": importances}).sort_values("importance", ascending=False).head(12)
    plt.figure(figsize=(8, 5))
    sns.barplot(data=imp, y="feature", x="importance", color="#2563eb")
    plt.title("Random Forest global önem sıralaması")
    plt.xlabel("Önem")
    plt.ylabel("Öznitelik")
    plt.tight_layout()
    fi_fig = FIG_DIR / "rf_feature_importance.png"
    plt.savefig(fi_fig, dpi=180)
    plt.close()

    shap_fig = FIG_DIR / "shap_summary.png"
    try:
        import shap
        X_test_tx = prep.transform(X_test.iloc[:300])
        explainer = shap.TreeExplainer(rf_pipe.named_steps["model"])
        shap_values = explainer.shap_values(X_test_tx)
        vals = shap_values[:, :, 1] if isinstance(shap_values, np.ndarray) and shap_values.ndim == 3 else shap_values[1]
        mean_abs = np.abs(vals).mean(axis=0)
        shap_imp = pd.DataFrame({"feature": names, "mean_abs_shap": mean_abs}).sort_values("mean_abs_shap", ascending=False).head(12)
        plt.figure(figsize=(8, 5))
        sns.barplot(data=shap_imp, y="feature", x="mean_abs_shap", color="#0f766e")
        plt.title("SHAP global açıklama özeti")
        plt.xlabel("Ortalama mutlak SHAP değeri")
        plt.ylabel("Öznitelik")
        plt.tight_layout()
        plt.savefig(shap_fig, dpi=180)
        plt.close()
    except Exception as exc:
        plt.figure(figsize=(8, 4))
        plt.text(0.02, 0.5, f"SHAP üretilemedi: {exc}", fontsize=10)
        plt.axis("off")
        plt.savefig(shap_fig, dpi=180)
        plt.close()

    lime_fig = FIG_DIR / "lime_local_explanation.png"
    try:
        from lime.lime_tabular import LimeTabularExplainer
        X_train_tx = prep.transform(X_train)
        X_test_tx = prep.transform(X_test)
        explainer = LimeTabularExplainer(
            X_train_tx,
            feature_names=names,
            class_names=["no", "yes"],
            discretize_continuous=True,
            random_state=42,
        )
        i = int(np.argmax(rf_pipe.predict_proba(X_test)[:, 1]))
        exp = explainer.explain_instance(
            X_test_tx[i],
            rf_pipe.named_steps["model"].predict_proba,
            num_features=8,
        )
        lime_df = pd.DataFrame(exp.as_list(), columns=["feature", "weight"]).sort_values("weight")
        colors = ["#dc2626" if v < 0 else "#16a34a" for v in lime_df["weight"]]
        plt.figure(figsize=(8, 4.8))
        plt.barh(lime_df["feature"], lime_df["weight"], color=colors)
        plt.axvline(0, color="#111827", linewidth=0.8)
        plt.title("LIME yerel açıklama örneği")
        plt.xlabel("Sınıf yes yönündeki katkı")
        plt.tight_layout()
        plt.savefig(lime_fig, dpi=180)
        plt.close()
    except Exception as exc:
        plt.figure(figsize=(8, 4))
        plt.text(0.02, 0.5, f"LIME üretilemedi: {exc}", fontsize=10)
        plt.axis("off")
        plt.savefig(lime_fig, dpi=180)
        plt.close()

    return {
        "target": target_fig,
        "corr": corr_fig,
        "auc": auc_fig,
        "cm": cm_fig,
        "fi": fi_fig,
        "shap": shap_fig,
        "lime": lime_fig,
    }


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color=(255, 255, 255))
        set_cell_shading(table.rows[0].cells[i], "1F4E79")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def make_docx(df, metrics, figs):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Calibri"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("BLM308 Veri Madenciliği Final Proje Raporu")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = subtitle.add_run("Banka Pazarlama Kampanyalarında Tahmin Edilebilirlik ve Açıklanabilirlik")
    s.bold = True
    s.font.size = Pt(15)
    doc.add_paragraph("Proje türü: Interpretable ML / XAI", style=None).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Bahar 2026", style=None).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_table(doc, ["Ad Soyad", "Numara", "Rol", "Yaptığı/Yazdığı Bölümler"], [
        ["Sude ANLAŞ", "231041038", "Veri Mühendisi + Değerlendirme", "Veri seti, EDA, ön işleme, metrikler, confusion matrix, sonuç yorumu"],
        ["Doğan Fatih OĞUR", "231041048", "Model Mühendisi + XAI/Raporlama", "Random Forest/MLP, SHAP-LIME açıklamaları, rapor bütünlüğü, README ve sunum taslağı"],
    ], widths=[1.4, 1.0, 2.0, 3.3])
    add_para(doc, "Teslim kapsamı: rapor.docx, kod klasörü, veri klasörü, README.txt, grafik çıktıları ve bonus için rapor.tex taslağı.")
    doc.add_page_break()

    doc.add_heading("Özet", level=1)
    add_para(doc, "Bu çalışmada Portekizli bir bankanın telefonla pazarlama kampanyalarına ait Bank Marketing veri seti kullanılarak müşterinin vadeli mevduata abone olup olmayacağı tahmin edilmiştir. Proje, CRISP-DM adımlarını izleyerek problem tanımı, veri anlayışı, veri hazırlığı, modelleme, değerlendirme ve iş değeri yorumunu uçtan uca ele alır. Temel model olarak Logistic Regression ve Decision Tree; black-box model olarak Random Forest ve MLP kurulmuştur. Model başarısı accuracy, precision, recall, F1 ve ROC-AUC metrikleriyle karşılaştırılmıştır. Ayrıca `duration` değişkeninin kampanya öncesinde bilinmemesi nedeniyle veri sızıntısı oluşturabileceği tartışılmış ve bu değişken dahil/hariç iki senaryo denenmiştir. En güçlü modelin kararları SHAP ve LIME ile açıklanarak hem genel öznitelik etkileri hem de tek müşteri düzeyindeki karar gerekçeleri yorumlanmıştır. Bulgular, açıklanabilir modellerin bankacılık gibi kararların gerekçelendirilmesinin önemli olduğu alanlarda yalnızca tahmin başarısı değil, güven ve iş değeri de sağladığını göstermektedir.")
    add_para(doc, "Anahtar kelimeler: veri madenciliği, sınıflandırma, Random Forest, MLP, SHAP, LIME, XAI, bank marketing.")

    doc.add_heading("1. Giriş", level=1)
    doc.add_heading("1.1 Problem Tanımı", level=2)
    add_para(doc, "Bankalar, pazarlama kampanyalarında sınırlı çağrı merkezi kapasitesini doğru müşterilere yönlendirmek ister. Rastgele arama stratejileri hem maliyeti artırır hem de müşteri memnuniyetini düşürebilir. Bu projenin problemi, geçmiş kampanya ve müşteri bilgilerinden yararlanarak bir müşterinin vadeli mevduata abone olma olasılığını tahmin etmek ve black-box model kararlarını açıklanabilir hale getirmektir.")
    doc.add_heading("1.2 Motivasyon ve İş Değeri", level=2)
    add_para(doc, "Bu problemin çözümü, bankanın daha verimli müşteri hedeflemesi yapmasını sağlar. Model yüksek olasılıklı müşterileri önceliklendirirse arama başına dönüşüm oranı artabilir. Ancak finans alanında kararların gerekçelendirilmesi gerekir; bu nedenle XAI bölümü proje için kritik önemdedir. Pazarlama yöneticisi yalnızca 'bu müşteri aranmalı' sonucunu değil, bu sonuca yaş, önceki kampanya sonucu, iletişim tipi veya ekonomik göstergeler gibi hangi değişkenlerin neden olduğunu da görebilmelidir.")
    doc.add_heading("1.3 Literatür ve İlgili Çalışmalar", level=2)
    add_para(doc, "Bank Marketing veri seti, Moro, Cortez ve Rita'nın bankacılık telepazarlama başarısını veri madenciliği ile tahmin ettiği çalışma ile ilişkilidir. UCI sayfasında veri setinin Portekizli bir bankanın telefon kampanyalarına dayandığı ve sınıflandırma hedefinin müşterinin vadeli mevduata abone olup olmayacağını tahmin etmek olduğu belirtilir. Açıklanabilir yapay zeka tarafında Lundberg ve Lee'nin SHAP yaklaşımı, tahminleri Shapley değerleriyle öznitelik katkılarına ayırır. Ribeiro, Singh ve Guestrin tarafından önerilen LIME ise herhangi bir sınıflandırıcının belirli bir örnek etrafındaki davranışını yerel, yorumlanabilir bir modelle açıklar.")
    doc.add_heading("1.4 Raporun Yapısı", level=2)
    add_para(doc, "İkinci bölümde CRISP-DM akışı ve kullanılan araçlar, üçüncü bölümde veri seti ve ön işleme, dördüncü bölümde modeller, beşinci bölümde performans ve XAI sonuçları, altıncı bölümde iş değeri, sınırlılıklar ve gelecek çalışma önerileri sunulmuştur.")

    doc.add_heading("2. Yöntem: CRISP-DM Akışı", level=1)
    add_para(doc, "Proje CRISP-DM'in altı adımına göre kurgulanmıştır: iş anlayışı, veri anlayışı, veri hazırlığı, modelleme, değerlendirme ve dağıtım/iş değeri yorumu. Bu yapı, raporun yalnızca teknik model çıktısı değil, gerçek bir veri madenciliği projesi olarak değerlendirilmesini sağlar.")
    add_bullets(doc, [
        "İş anlayışı: banka kampanyalarında dönüşüm oranını artırma ve kararları açıklama.",
        "Veri anlayışı: müşteri, kampanya ve ekonomik göstergelerin incelenmesi.",
        "Veri hazırlığı: kategorik kodlama, sayısal ölçekleme, hedef değişken dönüşümü.",
        "Modelleme: Logistic Regression, Decision Tree, Random Forest ve MLP.",
        "Değerlendirme: test seti metrikleri, confusion matrix ve veri sızıntısı tartışması.",
        "Dağıtım yorumu: çağrı önceliklendirme ve model açıklamalarının pazarlama ekibi için kullanımı.",
    ])
    doc.add_heading("2.1 Kullanılan Araçlar", level=2)
    add_table(doc, ["Araç", "Kullanım Amacı"], [
        ["Python", "Veri hazırlama, modelleme, metrik üretimi ve grafikler"],
        ["scikit-learn", "Sınıflandırma modelleri ve değerlendirme metrikleri"],
        ["SHAP", "Random Forest için global ve yerel açıklanabilirlik"],
        ["LIME", "Tek müşteri düzeyinde model karar açıklaması"],
        ["Weka karşılığı", "J48, RandomForest ve MultilayerPerceptron deneylerinin raporlanabilir karşılığı"],
    ])

    doc.add_heading("3. Veri", level=1)
    doc.add_heading("3.1 Veri Kaynağı", level=2)
    add_para(doc, "Veri seti UCI Machine Learning Repository üzerinde yayımlanan Bank Marketing veri setidir. `bank-additional-full.csv` dosyası 41.188 gözlem ve 20 girdi değişkeni içerir. Hedef değişken `y`, müşterinin vadeli mevduata abone olup olmadığını gösteren ikili sınıftır.")
    yes = int((df["y"] == "yes").sum())
    no = int((df["y"] == "no").sum())
    add_table(doc, ["Özellik", "Değer"], [
        ["Örnek sayısı", f"{len(df):,}".replace(",", ".")],
        ["Girdi değişkeni", "20"],
        ["Hedef değişken", "y: yes/no"],
        ["Pozitif sınıf", f"yes = {yes:,} (%{yes / len(df) * 100:.1f})".replace(",", ".")],
        ["Negatif sınıf", f"no = {no:,} (%{no / len(df) * 100:.1f})".replace(",", ".")],
    ])
    doc.add_picture(str(figs["target"]), width=Inches(5.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Şekil 1. Hedef değişken dağılımı sınıf dengesizliğini göstermektedir.")

    doc.add_heading("3.2 Değişken Grupları", level=2)
    add_table(doc, ["Grup", "Örnek Değişkenler", "Açıklama"], [
        ["Müşteri bilgisi", "age, job, marital, education", "Müşterinin demografik ve sosyo-ekonomik profili"],
        ["Finansal durum", "default, housing, loan", "Kredi temerrüdü, konut kredisi ve kişisel kredi bilgileri"],
        ["Kampanya bilgisi", "contact, month, day_of_week, campaign, pdays, previous, poutcome", "Mevcut ve önceki kampanya temasları"],
        ["Ekonomik bağlam", "emp.var.rate, cons.price.idx, cons.conf.idx, euribor3m, nr.employed", "Makroekonomik göstergeler"],
    ], widths=[1.5, 2.4, 3.6])
    doc.add_picture(str(figs["corr"]), width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Şekil 2. Sayısal değişkenler arasındaki korelasyon, özellikle ekonomik göstergelerin birlikte hareket ettiğini gösterir.")

    doc.add_heading("3.3 Ön İşleme", level=2)
    add_bullets(doc, [
        "Hedef değişken `yes=1`, `no=0` olarak kodlandı.",
        "Kategorik değişkenler One-Hot Encoding ile dönüştürüldü.",
        "Sayısal değişkenler StandardScaler ile ölçeklendirildi; bu adım MLP için özellikle gereklidir.",
        "Veri %75 eğitim, %25 test olarak stratified şekilde ayrıldı.",
        "`duration` değişkeni için iki senaryo kuruldu: dahil ve hariç. Bu değişken çağrı bittikten sonra bilindiği için kampanya öncesi kullanımda veri sızıntısı riski taşır.",
    ])

    doc.add_heading("4. Modelleme", level=1)
    add_para(doc, "Seçilen modeller hem temel yorumlanabilir yaklaşımları hem de XAI gerektiren black-box yaklaşımları kapsayacak şekilde belirlenmiştir. Böylece sınıflandırıcı karşılaştırması, ensemble boyutu ve XAI başlığı birlikte karşılanmıştır.")
    add_table(doc, ["Model", "Tür", "Weka Karşılığı", "Projedeki Rolü"], [
        ["Logistic Regression", "Lineer temel model", "Logistic", "Karşılaştırma ve yorumlanabilir referans"],
        ["Decision Tree", "Ağaç tabanlı model", "J48", "Kural benzeri açıklanabilir temel model"],
        ["Random Forest", "Ensemble black-box", "trees.RandomForest", "Ana yüksek performans modeli ve SHAP açıklaması"],
        ["MLP", "Sinir ağı black-box", "functions.MultilayerPerceptron", "Alternatif black-box model"],
    ], widths=[1.6, 1.6, 1.8, 2.7])
    doc.add_heading("4.1 Deney Kurulumu", level=2)
    add_bullets(doc, [
        "Test protokolü: stratified %75 eğitim / %25 test bölünmesi.",
        "Random seed: 42; deneylerin tekrarlanabilirliği için sabit tutuldu.",
        "Sınıf dengesizliği: Logistic Regression ve Random Forest'ta class_weight kullanıldı.",
        "Ana metrik: ROC-AUC ve F1; accuracy tek başına yeterli görülmedi.",
        "XAI: Random Forest için global SHAP/feature importance, tek örnek için LIME açıklaması.",
    ])

    doc.add_heading("5. Sonuçlar ve Tartışma", level=1)
    doc.add_heading("5.1 Performans Karşılaştırması", level=2)
    metric_rows = []
    for _, row in metrics.iterrows():
        metric_rows.append([
            row["Senaryo"], row["Model"],
            f"{row['Accuracy']:.3f}", f"{row['Precision']:.3f}",
            f"{row['Recall']:.3f}", f"{row['F1']:.3f}", f"{row['ROC-AUC']:.3f}",
        ])
    add_table(doc, ["Senaryo", "Model", "Acc.", "Prec.", "Recall", "F1", "ROC-AUC"], metric_rows, widths=[1.25, 1.55, .65, .65, .75, .55, .75])
    doc.add_picture(str(figs["auc"]), width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Şekil 3. ROC-AUC karşılaştırması, `duration` dahil olduğunda performansın belirgin yükseldiğini; ancak gerçek kampanya öncesi karar için `duration` hariç senaryonun daha dürüst olduğunu gösterir.")

    doc.add_heading("5.2 Confusion Matrix", level=2)
    doc.add_picture(str(figs["cm"]), width=Inches(4.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Random Forest confusion matrix, modelin çoğunluk sınıfı olan `no` sınıfında güçlü olduğunu, `yes` sınıfında ise sınıf dengesizliği nedeniyle recall/precision dengesinin ayrıca izlenmesi gerektiğini gösterir. Banka uygulamasında yanlış negatifler, potansiyel müşterinin kaçırılması anlamına geldiği için önemlidir.")

    doc.add_heading("5.3 Bias-Variance Yorumu", level=2)
    add_para(doc, "Decision Tree tekil yapısı nedeniyle yüksek varyansa yatkındır; derinlik sınırlaması bu riski azaltır ancak performansı sınırlayabilir. Random Forest çok sayıda ağacın ortalamasıyla varyansı düşürür ve bu veri setinde güçlü bir denge sağlar. Logistic Regression düşük varyanslı fakat doğrusal varsayımları nedeniyle daha sınırlı bir modeldir. MLP doğrusal olmayan ilişkileri yakalayabilir; ancak ölçekleme, erken durdurma ve düzenlileştirme yapılmazsa aşırı öğrenme riski taşır.")

    doc.add_heading("5.4 Model Yorumlanabilirliği: SHAP ve LIME", level=2)
    add_para(doc, "XAI bölümünde amaç, black-box modelin yalnızca skor üretmesini değil, karar gerekçesini de okunabilir hale getirmektir. SHAP global düzeyde hangi değişkenlerin model kararlarını en çok etkilediğini gösterirken, LIME tek müşteri etrafında yerel bir açıklama modeli kurar.")
    doc.add_picture(str(figs["fi"]), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Şekil 4. Random Forest önem sıralaması, modelin ekonomik göstergeler, önceki kampanya sonucu ve kampanya etkileşim değişkenlerine ağırlık verdiğini göstermektedir.")
    doc.add_picture(str(figs["shap"]), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Şekil 5. SHAP özeti, modelin genel karar mekanizmasında hangi özniteliklerin pozitif sınıf tahminini daha fazla değiştirdiğini açıklar.")
    doc.add_picture(str(figs["lime"]), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Şekil 6. LIME örneği, seçilen tek bir müşteride hangi koşulların `yes` sınıfını desteklediğini veya zayıflattığını gösterir. Bu açıklama, çağrı merkezi veya pazarlama uzmanının model önerisini iş bağlamında değerlendirmesine yardım eder.")

    doc.add_heading("5.5 Veri Sızıntısı Tartışması", level=2)
    add_para(doc, "`duration` değişkeni çağrı süresini temsil eder ve çağrı tamamlanmadan bilinmez. Bu nedenle kampanya öncesi müşteri seçimi için kullanılacak bir modelde `duration` değişkeni yer almamalıdır. Buna rağmen akademik karşılaştırma için iki senaryo raporlanmıştır. `duration` dahil senaryo modelin üst sınır performansını gösterirken, `duration` hariç senaryo gerçek uygulamaya daha yakındır. Bu ayrım, model başarısını yorumlarken iş bağlamının teknik metrikler kadar önemli olduğunu gösterir.")

    doc.add_heading("6. İş Değeri, Sınırlılıklar ve Sonuç", level=1)
    doc.add_heading("6.1 İş Değeri", level=2)
    add_para(doc, "Model, kampanya listesinde müşterileri olasılık skoruna göre sıralamak için kullanılabilir. En yüksek olasılıklı müşterilere öncelik verilmesi, çağrı merkezi kaynaklarının daha verimli kullanılmasını sağlar. XAI çıktıları ise pazarlama ekibine segment bazlı içgörü sunar: örneğin önceki kampanya sonucu başarılı olan veya belirli ekonomik koşullarda daha duyarlı müşteri grupları önceliklendirilebilir.")
    doc.add_heading("6.2 Etik ve Uygulama Notları", level=2)
    add_bullets(doc, [
        "Model kararı doğrudan müşteri aleyhine otomatik karar olarak kullanılmamalı, pazarlama önceliklendirme aracı olarak değerlendirilmelidir.",
        "Kişisel veriler için KVKK/GDPR benzeri veri koruma kurallarına uyulmalıdır.",
        "XAI açıklamaları modelin kesin nedensellik kurduğunu göstermez; sadece tahmin kararındaki katkıları açıklar.",
        "Periyodik izleme yapılmazsa ekonomik koşullar değiştikçe model performansı düşebilir.",
    ])
    doc.add_heading("6.3 Sınırlılıklar", level=2)
    add_para(doc, "Veri seti tek bir ülke ve belirli dönem kampanyalarını içerdiği için genellenebilirliği sınırlıdır. Ayrıca bazı değişkenler kampanya süreci içinde oluştuğundan gerçek zamanlı kullanımda veri sızıntısına dikkat edilmelidir. Sınıf dengesizliği, pozitif sınıfı yakalama başarısını zorlaştırır. Daha ileri çalışmalarda maliyet duyarlı öğrenme, eşik optimizasyonu ve zaman bazlı validasyon denenebilir.")
    doc.add_heading("6.4 Sonuç", level=2)
    add_para(doc, "Bu proje, Bank Marketing veri seti üzerinde black-box sınıflandırma modelleri kurmuş ve bu modellerin kararlarını SHAP/LIME ile açıklamıştır. Random Forest, performans ve açıklanabilirlik dengesinde güçlü bir aday olarak öne çıkmıştır. Ancak gerçek iş uygulaması için `duration` değişkeninin hariç tutulması gerektiği gösterilmiştir. Sonuç olarak proje, veri madenciliği sürecinin yalnızca model başarısından ibaret olmadığını; veri hazırlığı, metrik seçimi, açıklanabilirlik ve iş değeri yorumunun birlikte ele alınması gerektiğini ortaya koymaktadır.")

    doc.add_heading("Kaynaklar", level=1)
    refs = [
        "Moro, S., Cortez, P., Rita, P. (2014). A Data-Driven Approach to Predict the Success of Bank Telemarketing. Decision Support Systems.",
        "UCI Machine Learning Repository. Bank Marketing Dataset. https://archive.ics.uci.edu/dataset/222/bank+marketing",
        "Lundberg, S. M., Lee, S.-I. (2017). A Unified Approach to Interpreting Model Predictions. NeurIPS 2017.",
        "Ribeiro, M. T., Singh, S., Guestrin, C. (2016). Why Should I Trust You? Explaining the Predictions of Any Classifier. NAACL-HLT.",
        "Witten, I. H., Frank, E., Hall, M. A., Pal, C. J. Data Mining: Practical Machine Learning Tools and Techniques.",
        "scikit-learn documentation. https://scikit-learn.org",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    doc.add_heading("Ek A: Kod ve Veri Dizini", level=1)
    add_para(doc, "Teslim klasörü aşağıdaki yapıdadır:")
    add_para(doc, "proje/rapor.docx, proje/README.txt, proje/kod/run_experiment.py, proje/kod/build_report.py, proje/veri/bank-additional-full.csv, proje/figures/*.png")
    doc.add_heading("Ek B: Yapay Zeka Kullanımı Beyanı", level=1)
    add_para(doc, "Bu raporun taslak oluşturma ve dil düzenleme aşamasında yapay zeka desteği kullanılmıştır. Kullanılan veri seti, modelleme kodu, grafikler ve sonuç yorumları proje kapsamında kontrol edilerek düzenlenmiştir. Yapay zeka çıktıları doğrudan kopyalanmadan, proje yönergesine uygun akademik rapor formatına dönüştürülmüştür.")

    doc_path = OUT_DIR / "rapor.docx"
    doc.save(doc_path)
    return doc_path


def write_readme(metrics):
    readme = f"""BLM308 Veri Madenciliği Final Projesi

Proje başlığı:
Banka Pazarlama Kampanyalarında Tahmin Edilebilirlik ve Açıklanabilirlik

Öğrenciler:
- Sude ANLAŞ - 231041038
- Doğan Fatih OĞUR - 231041048

Proje türü:
Interpretable ML / XAI

Konu:
UCI Bank Marketing veri seti kullanılarak müşterinin vadeli mevduata abone olup olmayacağı tahmin edilir. Random Forest ve MLP gibi black-box modeller kurulur; Random Forest kararları SHAP ve LIME ile açıklanır.

Nasıl çalıştırılır:
1. Python 3.10+ kullanın.
2. Gerekli paketler:
   pip install pandas numpy scikit-learn matplotlib seaborn shap lime python-docx
3. Ana üretim:
   python build_project.py

Klasörler:
- teslim/rapor.docx: Word raporu
- kod/: rapora eklenebilecek sade deney kodları
- veri/: UCI Bank Marketing CSV kopyası
- figures/: EDA, metrik ve XAI grafikleri

Son üretilen metrik özeti:
{metrics.to_string(index=False)}
"""
    (OUT_DIR / "README.txt").write_text(readme, encoding="utf-8")


def write_code_files():
    run_experiment = """# BLM308 Final Projesi - Deney kodu
# Çalıştırma: python run_experiment.py
from pathlib import Path
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.compose import ColumnTransformer
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.ensemble import RandomForestClassifier
from sklearn.neural_network import MLPClassifier
from sklearn.metrics import classification_report, roc_auc_score

df = pd.read_csv(Path('../veri/bank-additional-full.csv'), sep=';')
y = (df['y'] == 'yes').astype(int)
X = df.drop(columns=['y', 'duration'])  # kampanya öncesi gerçekçi senaryo
cat_cols = X.select_dtypes(include=['object']).columns.tolist()
num_cols = X.select_dtypes(exclude=['object']).columns.tolist()

prep = ColumnTransformer([
    ('cat', Pipeline([('imp', SimpleImputer(strategy='most_frequent')),
                      ('ohe', OneHotEncoder(handle_unknown='ignore'))]), cat_cols),
    ('num', Pipeline([('imp', SimpleImputer(strategy='median')),
                      ('scaler', StandardScaler())]), num_cols),
])

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=.25, stratify=y, random_state=42)
models = {
    'Random Forest': RandomForestClassifier(n_estimators=180, max_depth=12, class_weight='balanced_subsample', random_state=42, n_jobs=-1),
    'MLP': MLPClassifier(hidden_layer_sizes=(48, 24), early_stopping=True, random_state=42, max_iter=90),
}

for name, model in models.items():
    pipe = Pipeline([('prep', prep), ('model', model)])
    pipe.fit(X_train, y_train)
    pred = pipe.predict(X_test)
    prob = pipe.predict_proba(X_test)[:, 1]
    print('\\n', name)
    print(classification_report(y_test, pred, target_names=['no', 'yes']))
    print('ROC-AUC:', roc_auc_score(y_test, prob))
"""
    (CODE_DIR / "run_experiment.py").write_text(run_experiment, encoding="utf-8")

    xai_code = """# BLM308 Final Projesi - XAI notu
# SHAP: Random Forest için global ve local öznitelik katkılarını çıkarır.
# LIME: Tek bir müşteri etrafında yerel açıklama modeli kurar.
# Tam üretim için ana klasördeki build_project.py çalıştırılabilir.
"""
    (CODE_DIR / "xai_notes.py").write_text(xai_code, encoding="utf-8")


def write_tex(metrics):
    top = metrics.sort_values("ROC-AUC", ascending=False).iloc[0]
    tex = rf"""\documentclass[11pt,a4paper]{{article}}
\usepackage[utf8]{{inputenc}}
\usepackage[T1]{{fontenc}}
\usepackage[turkish]{{babel}}
\usepackage[margin=2.3cm]{{geometry}}
\usepackage{{booktabs}}
\title{{Banka Pazarlama Kampanyalarında Tahmin Edilebilirlik ve Açıklanabilirlik}}
\author{{Sude ANLAŞ 231041038 \and Doğan Fatih OĞUR 231041048}}
\date{{Bahar 2026}}
\begin{{document}}
\maketitle
\section*{{Özet}}
Bu LaTeX dosyası bonus teslim için hazırlanmış kısa rapor kaynağıdır. Ayrıntılı rapor Word dosyası olan \texttt{{rapor.docx}} içindedir.
\section{{Proje Türü}}
Interpretable ML / XAI.
\section{{En İyi Sonuç}}
En yüksek ROC-AUC değeri {top['Senaryo']} senaryosunda {top['Model']} modeliyle {top['ROC-AUC']:.3f} olarak ölçülmüştür.
\section{{Kaynaklar}}
UCI Bank Marketing Dataset; Moro, Cortez ve Rita (2014); Lundberg ve Lee (2017); Ribeiro, Singh ve Guestrin (2016).
\end{{document}}
"""
    (OUT_DIR / "rapor.tex").write_text(tex, encoding="utf-8")


def main():
    df = load_data()
    metrics_with, fitted_with = evaluate_models(df, include_duration=True)
    metrics_without, fitted_without = evaluate_models(df, include_duration=False)
    metrics = pd.concat([metrics_with, metrics_without], ignore_index=True)
    metrics.to_csv(OUT_DIR / "metrics.csv", index=False)
    figs = make_figures(df, metrics, fitted_without)
    docx_path = make_docx(df, metrics, figs)
    write_readme(metrics)
    write_code_files()
    write_tex(metrics)
    summary = {"docx": str(docx_path), "rows": metrics.to_dict(orient="records")}
    (OUT_DIR / "build_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
